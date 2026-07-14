# -*- coding: utf-8 -*-
"""
Build / update the implementation report (实现文档.docx) on the Desktop.

Strategy:
- Load the existing template docx (keeps the cover page + logo image + chapter title).
- Delete the placeholder example body (the PetStore / 管理员模块 example) after the
  chapter title paragraph.
- Append the real project's module/function sections, each following the template's
  per-function format:
      功能描述 -> 模型代码 -> 模型代码测试 -> 控制器代码 -> 视图代码 -> 运行效果
- Code blocks are rendered with a monospace font (Courier New), size 9, noProof,
  matching the template's "HTML 预格式" code style intent.

Content for the modules is provided by `doc_content` (imported from _doc_content.py),
which the extraction agents / main flow populated.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:/Users/31815/Desktop/实现文档.docx"
OUT = r"C:/Users/31815/Desktop/实现文档.docx"   # overwrite in place

# Screenshots available in the repo (real app screenshots)
SHOT_JOIN = r"C:/Users/31815/IdeaProjects/task1/docs/verify-join-team-modal.png"
SHOT_AWARD = r"C:/Users/31815/IdeaProjects/task1/docs/verify-award-list.png"


def set_run_font(run, name="Courier New", size=9, bold=False, color=None):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # also set east-asian font so Chinese chars (rare in code) render mono
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rfonts.set(qn('w:eastAsia'), name)


def add_para(doc, text="", style=None, size=None, bold=False, align=None,
             font=None, color=None, space_after=4):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        if font:
            set_run_font(r, name=font, size=size or 10.5, bold=bold, color=color)
        else:
            if size:
                r.font.size = Pt(size)
            r.font.bold = bold
            if color is not None:
                r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _unescape(s):
    return s.replace("&gt;", ">").replace("&lt;", "<").replace("&amp;", "&")


def add_code_block(doc, code):
    """Add a code block (possibly multi-line) as a single paragraph, preserving
    line breaks, with monospace font and no proofing."""
    code = _unescape(code)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    # use HTML Preformatted style if present for fidelity
    try:
        p.style = "HTML 预格式"
    except Exception:
        pass
    # remove default tab stops / set left indent small
    lines = code.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=9, bold=False)
    return p


def add_subheading(doc, text, level):
    """level: 2 -> Heading 2 (module), 3 -> Heading 3 (function)"""
    style = "Heading %d" % level
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    # ensure east-asian font for the heading renders
    set_run_font(r, name="宋体", size=None)
    return p


def add_screenshot(doc, path, width_inch=5.5, caption=None):
    if not os.path.exists(path):
        add_para(doc, "（运行效果截图缺失：%s）" % path, size=9, color=RGBColor(0x88, 0x88, 0x88))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inch))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def render_function(doc, fn):
    """fn is a dict with keys: title, desc, model, test, ctrl, view, sql, shot, shot_caption, effect"""
    add_subheading(doc, fn["title"], 3)
    add_para(doc, fn["desc"], size=10.5, space_after=6)
    # 模型代码
    add_para(doc, "模型代码编写", bold=True, size=10.5, space_after=2)
    if fn.get("model"):
        add_code_block(doc, fn["model"])
    # 模型代码测试
    if fn.get("test"):
        add_para(doc, "模型代码测试", bold=True, size=10.5, space_after=2)
        add_code_block(doc, fn["test"])
    # 控制器代码
    if fn.get("ctrl"):
        add_para(doc, "控制器代码", bold=True, size=10.5, space_after=2)
        add_code_block(doc, fn["ctrl"])
    # 视图代码
    if fn.get("view"):
        add_para(doc, "视图代码（视图代码精选一部分）", bold=True, size=10.5, space_after=2)
        add_code_block(doc, fn["view"])
    # 关键SQL
    if fn.get("sql"):
        add_para(doc, "关键SQL", bold=True, size=10.5, space_after=2)
        add_code_block(doc, fn["sql"])
    # 运行效果
    add_para(doc, "运行效果", bold=True, size=10.5, space_after=2)
    if fn.get("shot"):
        add_screenshot(doc, fn["shot"], width_inch=fn.get("shot_w", 5.2),
                      caption=fn.get("shot_caption"))
    if fn.get("effect"):
        add_para(doc, fn["effect"], size=10.5, space_after=6)


def main():
    import shutil
    from _doc_content import MODULES, COVER_TEAM, COVER_MEMBERS, OVERVIEW

    # safety backup of the original template
    backup = SRC + ".bak"
    if not os.path.exists(backup):
        shutil.copy(SRC, backup)
        print("Backup created:", backup)

    doc = Document(SRC)

    # 1) Fill the cover team info (placeholders -> real info)
    #    Team name: left as-is unless COVER_TEAM provided.
    if COVER_TEAM:
        for p in doc.paragraphs:
            if p.text.startswith("团队名称"):
                _set_para_text(p, "团队名称  " + COVER_TEAM)
                break
    # fill member lines (two 团队成员 lines -> two strings)
    member_lines = list(COVER_MEMBERS)
    mi = 0
    for p in doc.paragraphs:
        if p.text.startswith("团队成员") and mi < len(member_lines):
            _set_para_text(p, "团队成员  " + member_lines[mi])
            mi += 1

    # 2) Clean the chapter title text (remove placeholder asterisks if any)
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            _set_para_text(p, "项目实现报告")
            break

    # 3) Delete everything from the first Heading 2 onward (the example body)
    #    Note: this template uses custom style IDs, so a Heading 2 paragraph's
    #    pStyle @w:val is "2" (not "Heading2"). Detect via python-docx style name.
    body = doc.element.body
    start_elem = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Heading 2":
            start_elem = p._element
            break
    if start_elem is not None:
        remove_list = []
        found = False
        for child in list(body):
            if child is start_elem:
                found = True
            if found and child.tag == qn('w:p'):
                remove_list.append(child)
        for child in remove_list:
            body.remove(child)

    # 4) Append overview + modules
    add_para(doc, OVERVIEW, size=10.5, space_after=8)

    for mod in MODULES:
        add_subheading(doc, mod["title"], 2)
        if mod.get("intro"):
            add_para(doc, mod["intro"], size=10.5, space_after=6)
        for fn in mod["functions"]:
            render_function(doc, fn)

    doc.save(OUT)
    print("Saved:", OUT)


def _set_para_text(p, text):
    """Replace all run text in a paragraph with `text` (keeps first run's formatting)."""
    elem = p._element
    runs = elem.findall(qn('w:r'))
    if not runs:
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        elem.append(r)
        return
    # clear all but first run text, set first run text
    first = True
    for r in runs:
        t = r.find(qn('w:t'))
        if t is None:
            t = r.makeelement(qn('w:t'), {})
            r.append(t)
        if first:
            t.text = text
            first = False
        else:
            elem.remove(r)


if __name__ == "__main__":
    main()
