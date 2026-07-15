from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "洪振博个人总结报告.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
ORANGE = "C65D21"
GREEN = "2E7D32"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color="CBD5E1", size="6"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="宋体", size=10.5, bold=False, color=DARK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, "等线", 9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, ORANGE),
    ):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("大学生海报设计竞赛系统  ·  洪振博个人总结")
    set_run_font(run, "等线", 9, color=MID_GRAY)
    set_paragraph_spacing(header, after=0, line=1.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("课程设计个人总结  |  ")
    set_run_font(run, "等线", 9, color=MID_GRAY)
    add_page_field(footer)
    set_paragraph_spacing(footer, after=0, line=1.0)


def add_body(doc, text, first_line=True, after=6, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, after=after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, "宋体", 10.5, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, "宋体", 10.5)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_note(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=LIGHT_BLUE, size="8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.35)
    r = p.add_run(label + "：")
    set_run_font(r, "黑体", 10.5, bold=True, color=BLUE)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.55)
    p.paragraph_format.first_line_indent = Cm(-0.38)
    set_paragraph_spacing(p, after=3, line=1.35)
    r = p.add_run("• " + text)
    set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = False
    return p


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_row = table.rows[0]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for idx, title in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, BLUE)
        set_cell_border(cell, color="FFFFFF", size="6")
        set_cell_margins(cell, top=110, start=100, bottom=110, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.15)
        r = p.add_run(str(title))
        set_run_font(r, "黑体", font_size, bold=True, color="FFFFFF")
        if widths and idx < len(widths):
            cell.width = Cm(widths[idx])

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 == 0 else GRAY)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and idx < len(widths):
                cell.width = Cm(widths[idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=0, line=1.2)
            r = p.add_run(str(value))
            set_run_font(r, "宋体", font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.4)
    set_paragraph_spacing(p, after=2, line=1.15)
    r = p.add_run(text)
    set_run_font(r, "Consolas", 9, color="334155")
    return p


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line=1.0)
    r = p.add_run("Web课程设计")
    set_run_font(r, "黑体", 18, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=20, line=1.0)
    r = p.add_run("个人总结报告")
    set_run_font(r, "黑体", 28, bold=True, color=BLUE)

    line = doc.add_table(rows=1, cols=1)
    line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = line.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_border(cell, color=BLUE, size="1")
    cell.height = Cm(0.16)
    cell.width = Cm(13.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=6, line=1.0)
    r = p.add_run("大学生海报设计竞赛系统")
    set_run_font(r, "黑体", 20, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28, line=1.0)
    r = p.add_run("—— 洪振博 ——")
    set_run_font(r, "宋体", 15, bold=True, color=ORANGE)

    info = [
        ("姓名", "洪振博"),
        ("项目角色", "副队长、竞赛管理模块负责人、基础设施与部署维护参与者"),
        ("学院 / 专业", "智能工程学院 / 计算机科学与技术"),
        ("项目周期", "2026年7月4日—2026年7月19日"),
        ("报告口径", "截至2026年7月15日的项目源码、文档与 Git 记录"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in info:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell, color="D5DEE8")
            set_cell_margins(cell, top=125, start=180, bottom=125, end=180)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], PALE_BLUE)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.15)
        r = p.add_run(label)
        set_run_font(r, "黑体", 10.5, bold=True, color=BLUE)
        p = cells[1].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.15)
        r = p.add_run(value)
        set_run_font(r, "宋体", 10.5)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=0, line=1.0)
    r = p.add_run("个人总结 · 工程实践 · 团队协作")
    set_run_font(r, "等线", 10, color=MID_GRAY)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "目录", 1)
    add_body(doc, "本报告围绕洪振博在大学生海报设计竞赛系统中的实际工作展开，既记录独立负责的竞赛管理与基础设施工作，也记录在跨模块联调、Bug 修复、服务器部署和数据库维护中的协作贡献。", first_line=False)
    rows = [
        ("一", "摘要", "—"),
        ("二", "项目背景与个人角色", "1"),
        ("三", "基础架构与数据库设计", "2"),
        ("四", "竞赛管理模块实现", "4"),
        ("五", "Bug 修复与稳定性治理", "6"),
        ("六", "服务器部署与数据库管理", "9"),
        ("七", "前端体验与进入页建设", "11"),
        ("八", "Git 协作与个人工作量证据", "13"),
        ("九", "团队协作、能力提升与个人反思", "15"),
        ("十", "不足、风险与后续计划", "17"),
        ("十一", "结语与证据索引", "19"),
    ]
    add_table(doc, ["序号", "章节", "页码参考"], rows, widths=[2, 9, 2.5], font_size=10)
    add_note(doc, "阅读说明", "页码为章节顺序参考，Word 打开后可根据实际排版更新页码。报告中引用的提交 ID 和文件路径用于复核个人贡献，不等同于对团队其他成员工作的否定。")
    doc.add_page_break()


def add_abstract(doc):
    add_heading(doc, "一、摘要", 1)
    add_body(doc, "本项目是一套面向大学生海报设计竞赛的线上管理系统，覆盖竞赛发布、队伍报名、作品提交、评委评分、评语管理、获奖设置、电子奖状和新闻公告等业务流程。洪振博在项目中承担副队长和竞赛管理模块负责人的职责，同时参与数据库设计、MVC 基础架构搭建、Tomcat 兼容性处理、跨模块联调、Bug 修复、服务器部署和远程数据库维护。")
    add_body(doc, "从 Git 历史看，截至 2026 年 7 月 15 日，洪振博署名提交共 70 次，其中排除合并提交后仍有 62 次非合并提交，时间覆盖 7 月 4 日基础建设至 7 月 14 日功能稳定性和测试补充。贡献内容既包括竞赛模块的完整实现，也包括角色首页、权限边界、图片访问、缩略图、获奖自动化、前端统一风格和测试代码等跨模块工作。")
    add_body(doc, "本总结不只罗列代码数量，而是从“问题—分析—方案—验证—遗留风险”的角度记录工作过程。报告中的服务器、数据库和测试结论以当前仓库配置、项目进展文档、构建产物、部署记录和 Git 证据为基础；对仍需在目标服务器上复测或尚未提交的内容，将明确标注其状态。")
    add_note(doc, "核心结论", "洪振博的主要价值在于把一个初始 MVC 骨架推进为能够持续联调、可以构建 WAR、具备远程数据库配置和多角色页面入口的完整工程，并在项目后期承担了较多“把系统跑起来、接起来、稳定下来”的工作。")


def add_background(doc):
    add_heading(doc, "二、项目背景与个人角色", 1)
    add_heading(doc, "2.1 项目目标与技术边界", 2)
    add_body(doc, "系统服务于竞赛组织者、评委、队长和队员四类角色。管理员需要发布竞赛、管理用户、设置获奖和发布公告；评委需要查看作品、打分和写评语；队长需要创建队伍、邀请或审核队员、报名竞赛并提交作品；队员需要注册登录、加入队伍、查看作品、点赞分享并查看奖状。")
    add_body(doc, "项目主系统采用 JSP、JavaScript、Bootstrap、Servlet、MySQL 8.0 和 MVC 分层架构，部署目标为使用 Java EE 规范的 Apache Tomcat 9.0。项目另有独立的 React 18 + Vite + GSAP 进入页，用于在用户进入正式系统前展示竞赛主题和作品氛围。")
    add_heading(doc, "2.2 洪振博的职责定位", 2)
    add_body(doc, "按照团队分工，洪振博是副队长，负责竞赛管理模块、技术架构、基础工具和项目实现材料。实际开发过程中，职责范围进一步扩展为：在项目初期搭建可并行开发的骨架；完成竞赛模块端到端链路；处理 Jakarta EE 与 Java EE、Jetty 本地运行与 Tomcat 目标部署之间的兼容问题；参与权限、图片、数据库和跨模块业务修复；负责将前端页面和服务器运行环境逐步接到统一的演示链路上。")
    add_table(doc, ["职责层次", "主要任务", "成果表现"], [
        ("模块负责人", "竞赛发布、分类、列表、详情、编辑、删除和状态逻辑", "DAO、Service、Servlet、JSP 与初始化数据形成完整链路"),
        ("技术骨干", "MVC 基础架构、Maven、数据库脚本、Java EE 适配", "为其他成员提供统一包结构、依赖和数据库边界"),
        ("联调维护者", "权限、作品、获奖、入队申请、图片访问等跨模块问题", "以日志、代码、页面和数据库字段联合定位问题"),
        ("部署维护者", "WAR 构建、Tomcat 配置、远程 MySQL、Navicat 脚本", "形成可部署到 8080 端口的项目产物和初始化路径"),
    ], widths=[3, 6, 5], font_size=9.2)


def add_architecture(doc):
    add_heading(doc, "三、基础架构与数据库设计", 1)
    add_heading(doc, "3.1 MVC 基础框架建设", 2)
    add_body(doc, "项目初始阶段需要先解决“所有人能不能同时开始写业务”的问题。洪振博参与完成了 Maven 依赖、数据库连接工具、模型类、DAO 接口和实现类、Service 接口和实现类、Servlet 控制器以及过滤器的骨架建设。该结构把页面渲染、请求分发、业务校验和 SQL 持久化分开，使竞赛、队伍、作品、评分和获奖等模块能够按职责并行推进。")
    add_body(doc, "在基础规范上，DAO 层统一使用 PreparedStatement，Service 层集中处理非空校验、默认值、角色和队长身份判断，Controller 层负责 action 路由与页面数据装载，工具类则承担数据库连接、密码处理、编码和文件处理。这样的分层不是为了增加文件数量，而是为了把多人开发时最容易发生的冲突点分开。")
    add_bullet(doc, "Maven 以 WAR 为打包目标，并纳入 MySQL Connector/J、JSTL、Commons FileUpload、Commons IO、Gson 和 JUnit 5 依赖。")
    add_bullet(doc, "源代码按 com.poster.controller、service、dao、model、util、filter 分包，便于模块负责人定位修改边界。")
    add_bullet(doc, "通过 EncodingFilter 统一处理请求编码，通过 AuthFilter 统一处理登录、公开资源和角色边界。")
    add_bullet(doc, "在早期 Jakarta 导入与 Tomcat 9 不兼容时，将 Servlet/JSTL/web.xml 相关配置切换到 javax Java EE 4.0 规范。")
    add_heading(doc, "3.2 数据库设计与版本演进", 2)
    add_body(doc, "数据库设计从竞赛主流程出发，建立用户与角色、竞赛与分类、队伍与成员、邀请与入队申请、作品与文件、评分与评语、获奖与奖状、新闻以及点赞分享等关系。项目进展文档早期记录为 17 张表；对当前 database/schema.sql 逐一核对后，脚本中实际包含 18 个 CREATE TABLE，差异主要来自后续补充的 team_application 入队申请表。报告采用“早期记录 17 张、当前脚本 18 张”的表述，以保持时间版本的准确性。")
    add_table(doc, ["数据域", "当前脚本中的核心表", "洪振博的相关工作"], [
        ("身份权限", "user、role、user_role", "参与基础实体与权限链路建设，配合 AuthFilter 和角色导航"),
        ("竞赛组织", "competition、competition_category", "完成竞赛 DAO、Service、Servlet、页面和初始化数据"),
        ("队伍流程", "team、team_member、invitation、team_application", "参与跨模块状态、报名、申请和权限问题联调"),
        ("作品交付", "work、work_file、work_like、work_share", "参与图片 BLOB、缩略图、访问策略和截止日期治理"),
        ("评审结果", "score、comment、award、certificate、news", "参与评分/获奖/奖状页面联调和一键获奖能力建设"),
    ], widths=[3, 6, 5], font_size=8.9)
    add_heading(doc, "3.3 数据库一致性意识", 2)
    add_body(doc, "在多人协作中，数据库字段变动比单个页面改动更容易造成连锁故障。洪振博参与了 schema、data、init_data 和 migrations 的同步检查，重点核对 user.avatar、work.image_data、work.image_content_type、thumbnail_data、thumbnail_content_type 等后续字段。V4 数据库一致性脚本还针对旧库字段、外键、状态含义和 Navicat 在 ONLY_FULL_GROUP_BY 下的执行问题提供了修复路径。")
    add_note(doc, "数据库版本口径", "全新环境优先执行 database/schema.sql，再执行 database/data.sql；已有旧库按具体版本选择 V2—V5 迁移脚本。init_data.sql 保留兼容旧文档的使用方式，但不能替代对目标库状态的核对。")


def add_competition(doc):
    add_heading(doc, "四、竞赛管理模块实现", 1)
    add_heading(doc, "4.1 后端三层链路", 2)
    add_body(doc, "竞赛模块是洪振博承担的核心业务模块。CompetitionDAOImpl 和 CategoryDAOImpl 负责竞赛及子类的增删改查，CompetitionServiceImpl 负责参数校验、默认状态和日期等业务规则，CompetitionServlet 负责 list、detail、add、edit、create、update、delete 等 action 的请求分发。通过这条链路，管理员可以创建竞赛、维护竞赛子类、查看详情、修改信息并删除无效数据。")
    add_body(doc, "在实现中，参数并不是从 JSP 直接拼接成 SQL，而是先由 Controller 接收，再由 Service 进行非空和合法性校验，最后由 DAO 通过 PreparedStatement 写入数据库。竞赛子类的动态新增、逐个删除和编辑同步也被纳入 CompetitionServlet 的保存流程，避免页面看似成功但分类数据没有落库。")
    add_heading(doc, "4.2 页面与业务状态联动", 2)
    add_body(doc, "竞赛列表、详情、发布和编辑页面共同组成管理端入口。洪振博后续又把竞赛详情页与队伍状态连接起来：未登录用户看到登录提示，已报名用户看到队伍与作品管理入口，尚未参赛用户看到创建队伍或搜索加入队伍入口。管理员专属的编辑和删除按钮则根据 Session 中的角色列表控制显示，减少普通用户误操作和越权入口。")
    add_table(doc, ["页面/控制器", "实现内容", "工程价值"], [
        ("competition_list.jsp", "列表、竞赛卡片、搜索/筛选和统计扩展", "让用户先了解赛事状态，再进入详情"),
        ("competition_add.jsp", "竞赛基本信息和动态子类输入", "减少硬编码分类，支持运营人员维护"),
        ("competition_edit.jsp", "已有子类展示、删除、新增和保存", "保证页面修改与关联表数据同步"),
        ("competition_detail.jsp", "详情、参赛状态、创建/加入队伍、管理按钮", "将竞赛、队伍和作品流程连接起来"),
        ("CompetitionServlet", "action 路由、统计查询、权限及状态装载", "承接 JSP 与 Service/DAO 之间的数据流"),
    ], widths=[4, 6, 4], font_size=9.1)
    add_heading(doc, "4.3 个人完成结果", 2)
    add_bullet(doc, "完成竞赛列表、详情、发布、编辑、删除和分类管理的端到端功能，项目进展记录中标注约 1200 行相关代码。")
    add_bullet(doc, "为管理员、评委和普通参赛用户设计差异化首页入口，管理员看到统计和管理卡片，评委看到评分工作台，队员看到竞赛与报名入口。")
    add_bullet(doc, "补充竞赛统计、作品列表筛选、分页数据结构 PageInfo 等后续能力，使模块从“能增删改查”向“可管理、可联调”推进。")


def add_bugs(doc):
    add_heading(doc, "五、Bug 修复与稳定性治理", 1)
    add_body(doc, "项目后期的工作重点从单个功能开发转向稳定性治理。洪振博参与的问题处理通常不是只修改一行页面代码，而是先确认问题属于依赖、路由、Session、数据库字段、文件路径还是业务状态，再同步修改控制器、Service、JSP 和 SQL。下面按问题类型记录代表性修复。")
    add_heading(doc, "5.1 运行环境与依赖兼容", 2)
    add_table(doc, ["问题", "定位与修复", "结果/经验"], [
        ("Tomcat 9 无法直接运行 Jakarta 代码", "将 Servlet/JSTL 依赖、web.xml 命名空间和 16 个 Java 文件中的 jakarta 导入切换到 javax；代表提交 0ad28e5", "统一到 Java EE 4.0/Tomcat 9 目标，降低团队环境差异"),
        ("首页入口与登录跳转不一致", "检查 IndexServlet、welcome-file、LandingServlet 和 AuthFilter 的路由/白名单，修复首页强制登录等入口问题", "使进入页、正式首页和登录/注册路径能够衔接"),
        ("本地运行与部署目标不一致", "保留 Jetty Maven 插件的本地 8080 配置，同时按 Tomcat 9 规范检查 WAR、web.xml 和静态资源放行", "本地验证与服务器部署有明确边界"),
    ], widths=[4, 7, 3], font_size=8.9)
    add_heading(doc, "5.2 权限、状态与 500/403 问题", 2)
    add_table(doc, ["问题", "修复动作", "验证关注点"], [
        ("普通用户可看到竞赛管理按钮", "在竞赛详情 JSP 中增加管理员角色判断，只对管理员显示编辑/删除入口", "不仅隐藏按钮，还要在 Servlet 入口再次校验"),
        ("管理员访问队伍/作品/申请路径出现 403", "调整 AuthFilter 的参与者和管理员判断，并统一导航入口；代表提交 2261a0b、fda9b76", "检查不同角色从首页、导航和直接 URL 访问的结果"),
        ("竞赛结束后仍显示提交作品按钮", "TeamServlet 传递 Competition，JSP 根据结束/取消状态关闭入口，WorkServlet 保留截止日期与状态双重校验", "前端提示和服务端拒绝必须同时存在"),
        ("JSP 页面运行时 500", "针对重复变量声明、实体方法名不一致和页面重复代码进行定位，参与修正 team_detail、award_manage 等页面", "以 JSP 编译后的 Servlet 视角检查变量与 getter"),
        ("获奖名单与当前比赛状态混淆", "将公开获奖列表调整为已结束竞赛的往届记录，并统一导航文案", "展示结果符合业务语义，避免未结束比赛提前进入公告"),
    ], widths=[4, 7, 3], font_size=8.7)
    add_heading(doc, "5.3 图片上传、访问与性能问题", 2)
    add_body(doc, "图片问题是跨本地文件、数据库、Servlet 输出和 JSP 访问路径的典型联调问题。早期方案尝试使用 work.image_data 的 MEDIUMBLOB 保存原图，并新增 ImageDataServlet 从数据库返回图片，解决了“图片只在上传者本机、其他设备无法显示”的问题；同时也发现所有查询都携带 BLOB 会导致 Navicat 和列表查询变慢。后续又增加缩略图字段和按场景选择原图/缩略图的访问方式：列表页使用缩略图，详情和评分页使用原图，缩略图缺失时回退到原图。")
    add_bullet(doc, "提交和更新作品时通过 ImageIO 生成最大宽度 300px、JPEG 质量约 0.75 的缩略图，减少列表页一次性加载原图的压力。")
    add_bullet(doc, "ImageDataServlet 支持 type=thumb/original，页面根据用途选择接口，避免把图片展示逻辑散落在各处。")
    add_bullet(doc, "文件上传工具校验扩展名、类型、大小并生成唯一文件名；外置存储目录避免 Tomcat 重新部署时丢失上传内容。")
    add_note(doc, "性能与数据遗留", "历史作品若没有原图 BLOB，就无法自动生成缩略图，需要重新上传或编辑后补写；当前工作区仍有上传/图片相关未提交修改，报告将其作为开发中的后续核验项。", fill="FFF4E8")
    add_heading(doc, "5.4 跨模块业务修复", 2)
    add_table(doc, ["场景", "处理内容", "体现的能力"], [
        ("同一竞赛重复建队", "在 Service 层查询用户在该竞赛中的有效队伍并拒绝重复创建", "把约束放在业务层，而不是只靠页面提示"),
        ("队伍人数上限", "由硬编码 5 改为读取 Competition.maxTeamSize，邀请和接受申请共用规则", "统一跨入口的业务规则"),
        ("搜索队伍并申请加入", "新增 DAO 模糊搜索、Service 方法、Servlet JSON action 和 AJAX Modal", "把正向邀请流程扩展为反向申请流程"),
        ("一键设置获奖", "补充 AutoAwardPolicy、候选作品筛选、名额计算、结果摘要和测试", "将重复人工操作抽象成可测试的业务策略"),
    ], widths=[4, 7, 3], font_size=8.9)
    add_heading(doc, "5.5 自动化测试核验结果", 2)
    add_body(doc, "2026 年 7 月 15 日使用 Maven Wrapper 执行项目现有 JUnit 5 测试，测试框架成功启动，共执行 47 个测试：44 个通过、3 个失败，没有出现 Error 或 Skipped。该结果说明项目已有权限、上传、路径、奖项策略、服务层和模板测试基础，但当前版本还不能写成“全部测试通过”。")
    add_table(doc, ["失败测试", "已定位的根因证据", "当前处理结论"], [
        ("CompetitionStatusPolicyTest.rejectsReopeningFinishedOrCancelledCompetition", "CompetitionStatusPolicy.java 允许 status 2→1；git 提交 d1423c3 在 7 月 14 日新增该分支，而测试仍要求 2→1 为 false", "属于状态策略与测试预期不一致，报告记录为待确认规则，不在本次总结中擅自修改"),
        ("ExhibitionUiTemplateTest.innerPageStylesAvoidKnownTemplateSlop", "app-pages.css 第 32 行仍包含 repeating-linear-gradient，命中测试禁止的样式模式", "属于统一视觉规范与现有样式残留不一致，需单独清理并复测"),
        ("UnifiedFrontendTemplateTest.everyNonParticipantJspLoadsTheSharedShell", "certificate_view.jsp 未引用 includes/app-shell-assets.jspf，其他页面的共享壳检查通过", "属于页面迁移遗漏，需补齐 include 或明确证书页的独立模板边界"),
    ], widths=[5.2, 5.8, 3], font_size=8.3)
    add_note(doc, "测试证据", "完整结果可在 target/surefire-reports 中复核。上述问题已完成错误信息、源码、测试和 Git 变更的交叉定位；本报告只记录根因和后续方向，不把未验证的修改写成已修复。", fill="FFF4E8")


def add_deployment(doc):
    add_heading(doc, "六、服务器部署与数据库管理", 1)
    add_heading(doc, "6.1 部署链路", 2)
    add_body(doc, "服务器部署不是把项目复制到 webapps 目录这么简单，需要同时保证构建产物、Servlet 规范、静态资源、上传目录和数据库连接一致。项目 pom.xml 的 packaging 为 war，当前工作区存在 target/task.war 构建产物；.smarttomcat/task/conf/server.xml 中配置了 8080 HTTP Connector。项目资料记录的演示地址为 http://120.26.46.0:8080/task/。")
    add_table(doc, ["环节", "实际配置/操作", "洪振博承担的工作"], [
        ("构建", "Maven WAR 打包；项目进展记录使用 mvnw -q -DskipTests package", "确认依赖、编译目标和 WAR 输出，处理构建失败与缺库问题"),
        ("容器", "Tomcat 9、Java EE 4.0、8080 端口；本地也保留 Jetty 9.4.58 运行插件", "完成 javax 兼容和 web.xml/Filter 检查"),
        ("应用", "部署 context 为 task，进入页由 index.html/LandingServlet 对接正式 JSP 路由", "处理 Landing、AuthFilter、静态资源和 JSP 入口衔接"),
        ("文件", "上传内容放到外置 storage 或 poster.storage.path 指定目录", "避免重新部署覆盖上传文件，并完善访问路径策略"),
        ("验证", "访问首页、登录、竞赛详情、队伍、作品、评分和获奖入口", "按角色和状态检查直接访问、导航和异常提示"),
    ], widths=[3, 7, 4], font_size=8.8)
    add_heading(doc, "6.2 远程 MySQL 服务器切换", 2)
    add_body(doc, "项目开发过程中数据库从旧远程主机切换到 120.26.46.0:3306 的 MySQL 8.0.46 环境，数据库名为 poster_competition，应用连接配置包含 SSL、时区、字符集和连接/读写超时参数。该项工作涉及 DBUtil.java、schema 脚本和初始化脚本的共同核对，不能只修改一个主机地址后就认为部署完成。")
    add_bullet(doc, "连接层使用 com.mysql.cj.jdbc.Driver，并设置 connectTimeout=5000、socketTimeout=10000，避免数据库不可达时请求无限等待。")
    add_bullet(doc, "新数据库初始化路径为先执行 database/schema.sql，再执行 database/data.sql；init_data.sql 用于兼容旧的初始化习惯。")
    add_bullet(doc, "已有旧数据库根据版本选择 V2__work_module.sql、V3__user_avatar.sql、V4__database_consistency.sql、V5__team_application.sql 等迁移脚本。")
    add_bullet(doc, "使用 Navicat 检查表结构、字段和执行结果，重点核对 user.avatar、work.image_data、work.image_content_type、thumbnail_data 和 thumbnail_content_type 等字段。")
    add_heading(doc, "6.3 数据库运维中的问题意识", 2)
    add_body(doc, "远程数据库管理让我认识到，表结构、初始化数据、迁移脚本和应用实体类必须版本同步。数据库管理不只是“建表”，还包括外键顺序、状态值含义、旧数据兼容、BLOB 查询性能和脚本可重复执行。V4 脚本针对旧库一致性和 Navicat/ONLY_FULL_GROUP_BY 场景提供了处理方式，体现了项目从本地开发向服务器环境迁移时的实际问题。")
    add_note(doc, "安全提醒", "当前 DBUtil.java 仍采用代码内连接配置。报告不记录密码；正式部署时应改为环境变量、Tomcat JNDI 或外部配置文件，并及时轮换已经出现在开发环境中的数据库凭据。", fill="FFF4E8")
    add_heading(doc, "6.4 建议的服务器复验清单", 2)
    add_table(doc, ["顺序", "复验动作", "通过标准"], [
        ("1", "检查 Java、Tomcat、MySQL 版本与端口", "Java 编译目标、Tomcat 9、MySQL 8.0.x、8080/3306 可达"),
        ("2", "新库执行 schema/data 或对应迁移脚本", "表、字段、外键、角色和测试账号完整"),
        ("3", "上传一张 JPG/PNG 作品并查看列表/详情", "其他浏览器或设备也能加载，缩略图/原图路径正确"),
        ("4", "按管理员、评委、队长、队员分别访问", "页面入口与 Servlet 权限判断一致"),
        ("5", "用结束竞赛数据走提交、评分、获奖流程", "截止状态、分数、奖状和公告数据关联正确"),
        ("6", "检查 Tomcat 日志和数据库慢查询", "无 JSP 500、连接泄漏、重复字段或大 BLOB 查询异常"),
    ], widths=[1.3, 8, 4.7], font_size=8.7)


def add_frontend(doc):
    add_heading(doc, "七、前端体验与进入页建设", 1)
    add_heading(doc, "7.1 角色化首页与统一导航", 2)
    add_body(doc, "洪振博参与重构 index.jsp 和 IndexServlet，使首页不再是所有用户看到相同的卡片集合，而是根据 Session 角色展示不同的工作入口。管理员看到竞赛、用户、新闻和统计管理；评委看到评分工作台和待评作品；队员看到竞赛大厅、报名、队伍和作品入口。随后又将公共导航抽取到 navbar.jspf，并在页面族样式中统一卡片、详情、表单、列表和工作台的视觉语言。")
    add_bullet(doc, "使用 app-shell、app-pages、role-home 等样式文件统一导航、阴影、卡片和页面间距。")
    add_bullet(doc, "清理多个页面中重复的紫色内联样式，使用页面族 class 管理统一风格。")
    add_bullet(doc, "通过 pointer-events 和 active 状态修复 Hero 轮播中不可见幻灯片覆盖链接的问题。")
    add_heading(doc, "7.2 React + GSAP Landing 进入页", 2)
    add_body(doc, "7 月 9—10 日，洪振博参与建设独立的 React 进入页，将项目从普通后台入口提升为具有竞赛主题的展示入口。页面包含 AnimatedText、MagneticLink、SpotlightCard 和 CinematicHero 等组件，使用现有 GSAP 实现滚动驱动视频、卡片 3D 变换、鼠标视差和 CTA 交互，并在 920px、560px 断点以及 prefers-reduced-motion 下做适配。")
    add_body(doc, "进入页和正式 JSP 系统之间通过 context path 对接 /index、/login、/register、/competition、/award 等入口。这个工作让我理解到，视觉展示不能破坏既有业务路由，前端增强必须同时考虑部署路径、AuthFilter 白名单、静态资源 MIME 类型和无动画降级。")
    add_heading(doc, "7.3 前端工作中的取舍", 2)
    add_body(doc, "页面改造并不是单纯增加动画。对参赛流程来说，重要的是用户能看到自己当前处于报名、组队、提交、评审还是获奖阶段。因此页面中的统计卡片、状态标签、步骤指示器、空状态和错误提示都服务于流程理解；动画则用于增强层次和作品展示感，并通过 reduced-motion 选项降低干扰。")


def add_git(doc):
    add_heading(doc, "八、Git 协作与个人工作量证据", 1)
    add_heading(doc, "8.1 提交记录概览", 2)
    add_body(doc, "当前 Git 仓库的 shortlog 显示洪振博署名提交 70 次；进一步排除 merge 提交后为 62 次非合并提交，包含功能、修复、文档、Revert/Reapply 和联调记录。提交次数不能直接代表质量，但能与文件变更、项目进度和源码内容互相印证，说明洪振博从基础架构到后期稳定性阶段持续参与。")
    add_table(doc, ["日期", "阶段重点", "代表提交/证据"], [
        ("7月4日", "基础框架、模型、DAO、Service、Servlet 和数据库起步", "9c887dd、8316252；早期 82 文件基础提交"),
        ("7月5日", "Tomcat/Java EE 适配、竞赛模块初版、首页入口修复", "0ad28e5、4a2ddfd、f8c8a30"),
        ("7月7—8日", "运行 Bug、依赖、数据库一致性、竞赛模块补充、图片方案", "dd69edb、07b3f02、ef9413a"),
        ("7月9—10日", "远程数据库切换、前端首页、React Landing 和视频交互", "573116c、f25eeaf、ea6e5bb"),
        ("7月11—12日", "权限边界、角色导航、结束比赛、申请加入队伍", "2261a0b、16409bf、fda9b76"),
        ("7月13日", "统一前端、稳定性补强、图片缩略图与测试", "4123ddb、831c54d、955b2ef"),
        ("7月14日", "自动获奖、分页/筛选、DAO 补充和 Service 测试", "70d0551、188e154、a8b1951"),
    ], widths=[2.2, 7.3, 4.5], font_size=8.8)
    add_heading(doc, "8.2 代表性提交解读", 2)
    add_table(doc, ["提交", "变更规模/内容", "个人贡献含义"], [
        ("9c887dd", "早期 82 个文件，搭建基础代码骨架", "完成从空项目到可分工开发的起点"),
        ("4a2ddfd", "竞赛模块 12 个文件，约 1046 行新增", "负责模块的端到端实现"),
        ("0ad28e5", "18 个文件，处理 jakarta→javax", "解决部署规范兼容性问题"),
        ("573116c", "前端基础与 Landing 资源，44 个文件", "将项目入口、展示层和正式系统接起来"),
        ("2261a0b", "33 个文件，统一角色导航和权限边界", "跨页面、Filter、Servlet 的系统性重构"),
        ("4123ddb", "81 个文件，稳定性、文件访问和测试补充", "从功能开发转向工程化质量治理"),
        ("70d0551", "自动获奖策略、服务、页面和测试", "把复杂人工流程抽象为可验证逻辑"),
        ("188e154", "18 个文件，分页、筛选和 DAO/Service 扩展", "继续完善管理端的数据处理能力"),
    ], widths=[2.1, 7.2, 4.7], font_size=8.7)
    add_note(doc, "工作量说明", "提交中包含合并、回滚和文档提交，因此不使用“总改动行数”作为个人贡献结论；报告优先使用职责、文件范围、提交主题和源码功能四类证据。")


def add_reflection(doc):
    add_heading(doc, "九、团队协作、能力提升与个人反思", 1)
    add_heading(doc, "9.1 团队协作", 2)
    add_body(doc, "项目采用垂直模块分工，每个成员负责一个相对完整的业务方向，但实际系统必须在数据库、Session、队伍状态、作品状态和角色权限之间联动。洪振博在承担竞赛模块的同时，持续参与合并、冲突处理、接口对齐和运行问题定位，将“自己的模块能运行”推进到“其他模块能够从竞赛入口接入”。")
    add_bullet(doc, "与队长协作处理基础框架、用户认证入口、Maven 依赖和 Java EE 兼容问题。")
    add_bullet(doc, "与队伍模块协作核对报名状态、队伍搜索、申请加入、人数上限和竞赛详情入口。")
    add_bullet(doc, "与作品模块协作处理作品详情、图片 BLOB、缩略图、截止日期和服务器文件路径。")
    add_bullet(doc, "与评分/获奖模块协作处理评分工作台、获奖列表、证书入口、已结束竞赛和一键获奖流程。")
    add_heading(doc, "9.2 能力提升", 2)
    add_table(doc, ["能力", "项目中的具体体现", "形成的认识"], [
        ("需求拆解", "把竞赛发布拆为分类、状态、详情、权限、统计和跨模块入口", "业务功能必须围绕状态和角色拆分"),
        ("后端工程", "Servlet action、Service 规则、DAO 参数化 SQL、Session 角色判断", "控制器不是业务规则的唯一位置"),
        ("数据库维护", "远程 MySQL 切换、schema/migration、BLOB 与缩略图、Navicat 执行", "数据库版本必须和代码一起管理"),
        ("部署排障", "WAR、Tomcat 9、8080、静态资源和外置上传目录", "本地通过不等于服务器可用"),
        ("质量意识", "安全策略类、JUnit 测试、状态校验和错误场景", "稳定性来自边界检查和可复现验证"),
        ("协作沟通", "Git 合并、冲突处理、进展文档和证据索引", "多人开发需要明确约定和记录"),
    ], widths=[3, 7, 4], font_size=8.8)
    add_heading(doc, "9.3 个人反思", 2)
    add_body(doc, "这次项目让我最大的收获是认识到，Java Web 项目的难点不只在于写出 CRUD。真正影响演示和交付的是环境兼容、数据库字段一致、权限边界、文件生命周期、状态流转和异常处理。竞赛详情页一个按钮是否显示，背后可能同时依赖 Session 角色、竞赛 status、队伍 status、截止日期和服务器数据库中的数据。")
    add_body(doc, "同时，前期一些提交信息较笼统、配置中存在敏感信息、部分功能在代码完成后还需要目标服务器复验，这些问题说明开发效率不能替代工程纪律。后续应进一步提高提交信息质量、把环境配置外置、建立完整的集成测试和部署检查记录。")


def add_risks(doc):
    add_heading(doc, "十、不足、风险与后续计划", 1)
    add_body(doc, "截至报告时间，项目已经具备较完整的功能骨架和演示链路，但仍存在需要在提交前继续核验的事项。以下内容不作为对个人工作的否定，而是对工程交付状态的如实记录。")
    add_table(doc, ["风险/不足", "当前状态", "后续计划"], [
        ("数据库凭据写在 DBUtil.java", "开发环境可连接远程库，但存在泄露和轮换风险", "改为环境变量/Tomcat JNDI/外部配置，轮换现有密码"),
        ("目标服务器完整回归", "已有 WAR、Tomcat 和公网访问记录，仍需按角色和完整流程复测", "固定部署版本，记录日志、数据库和页面验收结果"),
        ("图片历史数据与缩略图", "新上传路径已支持缩略图，旧数据可能没有 image_data", "补充数据迁移/重传策略，验证大图和缩略图性能"),
        ("工作区未提交修改", "上传、图片、作品详情等文件仍有本地改动", "先逐文件检查 diff，再完成测试、提交和部署"),
        ("JSP 输出安全", "项目进展记录曾指出部分页面仍需加强 XSS 输出转义", "将可变文本逐步改为 JSTL c:out 或统一 HTML 转义"),
        ("跨模块集成测试", "已有部分 JUnit 和模板测试，完整注册→报名→提交→评分→获奖仍需演练", "准备固定测试账号、测试竞赛和可重复数据脚本"),
        ("现有自动化测试未全绿", "2026-07-15 执行 47 个测试，44 通过、3 失败；失败原因已定位到状态规则、样式残留和证书页共享壳", "先确认 status 2→1 业务规则，再修复两个模板问题并重新运行完整测试"),
    ], widths=[4.2, 6.2, 3.6], font_size=8.6)
    add_heading(doc, "10.1 后续工作优先级", 2)
    add_bullet(doc, "第一优先级：保护远程数据库凭据，核对服务器实际部署版本，完成数据库备份和回滚方案。")
    add_bullet(doc, "第二优先级：在目标服务器执行完整角色流程和图片访问测试，重点观察 Tomcat 日志、数据库连接和 BLOB/缩略图查询性能。")
    add_bullet(doc, "第三优先级：整理当前未提交修改，补充边界测试和前端 XSS 输出处理，再按规范提交。")
    add_bullet(doc, "第四优先级：将部署步骤、数据库初始化、迁移顺序和测试账号整理成一份可交接文档，降低团队后续答辩演示风险。")


def add_conclusion(doc):
    add_heading(doc, "十一、结语与证据索引", 1)
    add_body(doc, "在大学生海报设计竞赛系统中，洪振博的工作主线可以概括为“搭框架、做模块、解问题、接环境、保联调”。从 7 月 4 日的数据库与 MVC 基础，到 7 月 5—8 日的竞赛管理，再到 7 月 9—14 日的远程数据库切换、进入页、角色权限、图片访问、缩略图、自动获奖和测试补强，工作范围随着系统复杂度增长而扩大。")
    add_body(doc, "个人总结不仅是对代码数量的汇报，更是对工程责任的说明。洪振博承担的价值在于把功能实现和运行环境连接起来：既关注 DAO 和 JSP 的细节，也关注 Tomcat、MySQL、WAR、文件存储、权限和部署后的可验证性。后续只要继续补齐安全配置、完整回归和交付文档，项目就能更稳定地完成最终答辩和提交。")
    add_heading(doc, "关键证据索引", 2)
    add_table(doc, ["证据类别", "文件/提交", "可核对内容"], [
        ("项目进展", "CLAUDE.md", "按日期记录个人负责模块、Bug 修复、部署和数据库事项"),
        ("基础架构", "pom.xml、src/main/webapp/WEB-INF/web.xml", "WAR、依赖、Java EE 4.0、Filter 和欢迎页配置"),
        ("竞赛模块", "CompetitionServlet.java、CompetitionServiceImpl.java、CompetitionDAOImpl.java", "竞赛与分类的 Controller/Service/DAO 链路"),
        ("数据库", "database/schema.sql、database/data.sql、database/migrations/", "表结构、初始化数据和旧库迁移路径"),
        ("部署", ".smarttomcat/task/conf/server.xml、target/task.war", "Tomcat 8080 配置与当前构建产物"),
        ("远程连接", "src/main/java/com/poster/util/DBUtil.java", "远程 MySQL 主机、数据库名、超时和驱动配置；密码未写入本报告"),
        ("文件与图片", "FileUploadUtil.java、ImageDataServlet.java、WorkServlet.java", "上传目录、图片访问、原图/缩略图和截止日期逻辑"),
        ("质量治理", "src/test/java/、AutoAwardPolicy.java", "权限、上传、状态、奖项名额和服务层测试"),
        ("Git", "9c887dd、0ad28e5、4a2ddfd、573116c、2261a0b、4123ddb、70d0551、188e154", "从基础、模块、部署、联调到稳定性治理的代表性提交"),
        ("团队材料", "C:/Users/Lenovo/Desktop/团队总结.docx", "系统定位、团队分工和公网演示地址背景"),
    ], widths=[3, 6.7, 4.3], font_size=8.5)
    add_note(doc, "最终说明", "本报告由项目源码、Git 历史、CLAUDE.md、服务器配置、数据库脚本和已有团队总结共同整理。对无法从当前证据确认的内容使用“参与、协作、待复验”等表述，不把计划、未提交改动或他人独立完成的工作全部归为洪振博个人成果。")


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_abstract(doc)
    add_background(doc)
    add_architecture(doc)
    add_competition(doc)
    add_bugs(doc)
    add_deployment(doc)
    add_frontend(doc)
    add_git(doc)
    add_reflection(doc)
    add_risks(doc)
    add_conclusion(doc)
    doc.core_properties.title = "洪振博个人总结报告"
    doc.core_properties.subject = "大学生海报设计竞赛系统 Web 课程设计"
    doc.core_properties.author = "洪振博"
    doc.core_properties.comments = "根据项目源码、Git 记录、服务器配置、数据库脚本和团队总结整理"
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
