# -*- coding: utf-8 -*-
"""Generate the team summary from the supplied Word template.

The source cover is kept as-is. Only the placeholder body after the cover's
page-break paragraph is replaced with the project summary content.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"D:\qiyeweixin\WXWork\1688854684669597\Cache\File\2026-07\团队总结.docx")
DEFAULT_OUTPUT = Path(r"C:\Users\Lenovo\Desktop\java\task\团队总结-已生成.docx")


TECHNICAL_PARAGRAPHS = [
    (
        "本项目围绕大学生海报设计竞赛的线上管理需求，建设了“大学生海报设计竞赛系统”。"
        "系统面向管理员、评委、队长和队员四类角色，覆盖竞赛发布、队伍报名、作品提交、评审以及获奖公布等业务环节，"
        "旨在把原本分散的竞赛管理流程集中到统一平台中。项目已部署到 http://120.26.46.0:8080/task/ 进行访问和展示。"
    ),
    (
        "系统采用 JSP + JavaScript + Bootstrap 作为前端技术，使用 Servlet + JSP 完成后端请求处理，"
        "以 MySQL 8.0 保存业务数据，并按照 Apache Tomcat 9.0 的 Java EE 规范进行部署。整体采用 MVC 分层架构，"
        "将页面展示、请求控制、业务服务和数据访问分离，便于多人并行开发、调试和后续维护。"
    ),
    (
        "数据库部分完成了竞赛、分类、用户、角色、队伍、成员、作品、评分、奖项和新闻等核心数据表设计，"
        "共形成 17 张表及对应的实体模型。DAO 层统一使用 PreparedStatement 访问数据库，Service 层负责输入校验、"
        "权限判断和跨表业务逻辑，Controller 层负责 Servlet 请求分发，形成了较为完整的三层业务链路。"
    ),
    (
        "目前项目已经完成用户注册登录、Session 权限校验、个人信息管理、竞赛列表与详情、竞赛发布与编辑、"
        "队伍创建、队长自动入队、队伍详情展示、成员移除、队伍搜索和统计概览等核心功能。队伍页面还完成了"
        "渐变封面、成员头像、队长标识、Tab 切换、搜索过滤和空状态等界面优化，提升了系统的可用性和视觉一致性。"
    ),
    (
        "在安全与工程规范方面，项目加入了密码加密存储、SQL 注入防护、登录过滤器、角色权限校验、统一编码处理和"
        "文件上传大小限制等措施。作品管理、点赞分享、评分获奖和新闻管理等模块已完成基础架构或进入后续开发阶段，"
        "团队将继续围绕完整流程联调、跨模块测试和答辩演示进行完善。"
    ),
]

TECHNICAL_POINTS = [
    "前后端分层：JSP 负责页面呈现，Servlet 负责请求控制，Service 负责业务规则，DAO 负责数据持久化。",
    "数据库与模型对应：17 张业务表均有相应实体类，便于在各模块之间传递结构化数据。",
    "权限与安全：登录状态由过滤器统一检查，敏感操作进行角色或队长身份验证，数据库查询使用参数化语句。",
    "可维护性：公共连接、编码、密码处理和文件处理能力集中在工具类中，减少重复实现。",
]

COOPERATION_PARAGRAPHS = [
    (
        "团队采用垂直功能模块分工，每位成员负责一个相对完整的业务模块，同时遵循统一的包结构、命名方式、"
        "MVC 分层和数据库访问规范。这样的分工让成员能够独立完成页面、Servlet、Service、DAO 以及相关数据表的开发，"
        "也降低了多人同时修改同一层代码造成的集成风险。"
    ),
    (
        "队长负责整体协调、用户认证模块和开发进度；副队长负责竞赛管理、技术架构和项目实现材料；杨祥博负责队伍管理"
        "以及界面体验优化；队员B负责作品管理和文件上传方向；队员C负责评分、获奖、新闻方向以及系统集成测试和问题修复。"
    ),
    (
        "团队通过每日晨会同步进度，汇报已完成内容、当天计划和遇到的问题；开发过程中使用功能分支和规范化提交信息，"
        "在模块完成后进行合并和联调。遇到数据库字段映射、Tomcat 规范适配、权限校验和页面交互等问题时，团队先定位边界，"
        "再通过代码审查、运行验证和共同讨论完成修正。"
    ),
    (
        "本次合作使团队成员不仅熟悉了 Java Web 项目的完整开发流程，也认识到需求拆分、接口约定、文档记录和集成测试"
        "对于多人协作的重要性。后续团队将继续以“注册登录—竞赛报名—作品提交—评分—获奖”的完整链路为主线，"
        "集中验证跨模块数据流和异常场景，保证系统在答辩演示时稳定运行。"
    ),
]

PERSONAL_SUMMARIES = [
    (
        "队长",
        "负责项目整体协调、进度推进和基础框架建设，完成用户注册登录、Session 管理、权限过滤和个人信息功能。"
        "在协作过程中负责统一技术规范、组织模块联调并跟进集成问题，对系统从需求分析到部署运行的全过程有了更完整的认识。",
    ),
    (
        "副队长",
        "负责竞赛管理模块和项目技术架构，完成竞赛列表、详情、发布、编辑、分类及相关数据访问功能。"
        "同时参与数据库设计、Tomcat 规范适配和实现报告整理，体会到业务字段设计、输入校验与页面交互需要同步考虑。",
    ),
    (
        "杨祥博",
        "负责队伍管理模块，完成创建队伍、队长自动入队、队伍详情、成员管理、移除成员和队伍展示优化等工作。"
        "通过对队伍卡片、成员头像、统计概览、Tab 页面和搜索过滤进行设计，进一步提升了前端页面组织和用户体验优化能力。",
    ),
    (
        "队员B",
        "负责作品管理方向，围绕作品提交、文件上传、作品展示、修改删除、点赞分享和图片预览等需求完善模块。"
        "在后续开发中将重点关注上传文件校验、文件路径安全、作品状态流转以及作品模块与队伍和评分模块之间的数据衔接。",
    ),
    (
        "队员C",
        "负责评分与获奖方向，并承担系统集成测试和问题修复工作，围绕评委评分、评语、奖项、奖状和新闻发布梳理业务流程。"
        "后续将重点验证评分范围、重复评分、获奖关联、证书生成和完整流程演示，提升系统功能的稳定性和可验证性。",
    ),
]


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def format_body_paragraph(paragraph, first_line=True, space_after=6):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.first_line_indent = Pt(-12)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("• " + text)
    set_run_font(run)
    return paragraph


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, name="黑体", size=16, bold=True, color=(31, 78, 121))
    return paragraph


def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, name="黑体", size=13, bold=True)
    return paragraph


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


def set_cell_border(cell, color="B7C9D6", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_photo_placeholder(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("以下位置用于插入团队合影")
    set_run_font(run, name="宋体", size=11, color=(100, 100, 100))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(5.8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)
    row = table.rows[0]
    row.height = Inches(3.2)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_paragraph.paragraph_format.space_after = Pt(0)
    run = cell_paragraph.add_run("[请在此处插入团队合影]")
    set_run_font(run, name="宋体", size=13, color=(127, 127, 127))


def remove_body_after_cover(doc, cover_end):
    body = doc.element.body
    after_cover = False
    for child in list(body):
        if child is cover_end:
            after_cover = True
            continue
        if after_cover and child.tag != qn("w:sectPr"):
            body.remove(child)


def locate_cover_end(doc):
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "一、技术总结":
            previous = paragraphs[index - 1] if index > 0 else None
            if previous is not None and 'w:type="page"' in previous._p.xml:
                return previous._p
    raise ValueError("未找到模板封面后的分页位置，停止生成以避免覆盖封面。")


def build_document(source: Path, output: Path):
    if not source.exists():
        raise FileNotFoundError(f"模板不存在：{source}")

    document = Document(str(source))
    paragraphs_before = list(document.paragraphs)
    cover_end = locate_cover_end(document)
    cover_text_before = [p.text for p in paragraphs_before if p._p is cover_end or p._p in list(cover_end.itersiblings(preceding=True))]
    cover_image_count = len(document.inline_shapes)

    remove_body_after_cover(document, cover_end)

    add_section_heading(document, "一、技术总结")
    for paragraph in TECHNICAL_PARAGRAPHS:
        add_body_paragraph(document, paragraph)
    for point in TECHNICAL_POINTS:
        add_bullet(document, point)

    add_page_break(document)
    add_section_heading(document, "二、团队合作总结")
    for paragraph in COOPERATION_PARAGRAPHS:
        add_body_paragraph(document, paragraph)

    add_page_break(document)
    add_section_heading(document, "三、团队合影")
    add_photo_placeholder(document)

    add_page_break(document)
    add_section_heading(document, "四、个人总结")
    for role, summary in PERSONAL_SUMMARIES:
        add_subheading(document, role)
        add_body_paragraph(document, summary)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))

    # The list is intentionally computed after saving only for a concise build log;
    # no cover paragraphs are edited by this generator.
    rebuilt = Document(str(output))
    cover_text_after = [p.text for p in rebuilt.paragraphs[: len(cover_text_before)]]
    if cover_image_count != len(rebuilt.inline_shapes):
        raise AssertionError("生成后模板图片数量发生变化。")
    if cover_text_before != cover_text_after:
        raise AssertionError("生成后封面文字发生变化。")

    print(f"source={source}")
    print(f"output={output}")
    print(f"paragraphs={len(rebuilt.paragraphs)}")
    print(f"inline_shapes={len(rebuilt.inline_shapes)}")
    print("cover=preserved")


def main():
    parser = argparse.ArgumentParser(description="Generate the project team summary Word document.")
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_document(args.source, args.output)


if __name__ == "__main__":
    main()
